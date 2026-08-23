"""Render the final research report from Markdown to a styled DOCX.

This script is intentionally independent of the project's econometric
environment. Run it with the isolated Codex document Python runtime, which
provides python-docx and Pillow.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4E79"
GREY = "59636E"
LIGHT_GREY = "F4F6F9"
TABLE_WIDTH_DXA = 9360


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = qn(f"w:{edge}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), BLUE)
    r_pr.append(colour)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    token_re = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|https?://\S+|`[^`]+`)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            run.bold = bold
            run.italic = italic
        token = match.group(0)
        if token.startswith("["):
            label, url = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            _add_hyperlink(paragraph, label, url)
        elif token.startswith("http"):
            trailing = ""
            while token and token[-1] in ".,;":
                trailing = token[-1] + trailing
                token = token[:-1]
            _add_hyperlink(paragraph, token, token)
            if trailing:
                paragraph.add_run(trailing)
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.2)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        run.bold = bold
        run.italic = italic


def _add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def _configure_document(doc: Document, meta: dict[str, str]) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 31, 31)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, before, after, colour in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Caption Report" not in doc.styles:
        caption = doc.styles.add_style("Caption Report", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Caption Report"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GREY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.0

    if "Reference" not in doc.styles:
        ref = doc.styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = doc.styles["Reference"]
    ref.font.name = "Calibri"
    ref.font.size = Pt(8.7)
    ref.paragraph_format.space_after = Pt(5)
    ref.paragraph_format.line_spacing = 1.0
    ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "Executive Summary Body" not in doc.styles:
        executive = doc.styles.add_style("Executive Summary Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        executive = doc.styles["Executive Summary Body"]
    executive.font.name = "Calibri"
    executive.font.size = Pt(9.5)
    executive.font.color.rgb = RGBColor(31, 31, 31)
    executive.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    executive.paragraph_format.space_before = Pt(0)
    executive.paragraph_format.space_after = Pt(5)
    executive.paragraph_format.line_spacing = 1.15

    header = sec.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    run = p.add_run("AU NEM RENEWABLES & PRICES")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GREY)
    p.add_run("\t")
    run = p.add_run("RESEARCH REPORT")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREY)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "B7C9DA")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)

    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Jixin Guo  •  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREY)
    _add_field(p, "PAGE")

    first_footer = sec.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Jixin Guo  •  Research report  •  23 August 2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREY)

    props = doc.core_properties
    props.title = meta["title"]
    props.author = meta["author"]
    props.subject = "Renewable penetration and wholesale electricity outcomes in Australia's NEM"
    props.keywords = "Australia; NEM; renewable energy; wholesale prices; negative prices; volatility"


def _add_cover(doc: Document, meta: dict[str, str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("APPLIED ENERGY ECONOMICS")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(meta["title"])
    run.font.name = "Calibri Light"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(meta["subtitle"])
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(GREY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(f"{meta['author']}  |  {meta['date']}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GREY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), BLUE)
    border.append(bottom)
    p_pr.append(border)


def _add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    _add_inline(p, text, bold=True)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EAF2F8")
    p_pr.append(shading)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    border.append(left)
    p_pr.append(border)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = [2100, 1900, 1400, 3960]
    _set_table_geometry(table, widths)
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 3) else WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(8.7)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
                _set_cell_shading(cell, LIGHT_GREY)
        if r_idx == 0:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
    p = doc.add_paragraph("Table 1. Headline fixed-effect estimates. Standard errors are clustered by AEST ISO week; N = 210,399.")
    p.style = doc.styles["Caption Report"]


def _set_picture_alt_text(paragraph, title: str, description: str) -> None:
    drawings = paragraph._p.xpath(".//wp:docPr")
    for node in drawings:
        node.set("title", title)
        node.set("descr", description)


def _add_figure(doc: Document, root: Path, spec: str) -> None:
    path_text, caption = [part.strip() for part in spec.split("::", 1)]
    image_path = root / path_text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    width = Inches(5.0) if image_path.name.startswith("fig1_") else Inches(6.45)
    p.add_run().add_picture(str(image_path), width=width)
    _set_picture_alt_text(p, caption.split(".", 1)[0], caption)
    cp = doc.add_paragraph(caption)
    cp.style = doc.styles["Caption Report"]


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    idx = start
    while idx < len(lines) and not lines[idx].strip():
        idx += 1
    rows: list[list[str]] = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        values = [v.strip() for v in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        idx += 1
    return rows, idx


def _add_page_content(
    doc: Document,
    root: Path,
    lines: list[str],
    *,
    default_style: str = "Normal",
) -> None:
    idx = 0
    in_references = False
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue
        if line.startswith("# "):
            doc.add_paragraph(line[2:], style="Heading 1")
            idx += 1
            continue
        if line.startswith("## "):
            title = line[3:]
            in_references = title == "References"
            doc.add_paragraph(title, style="Heading 2")
            idx += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 3")
            idx += 1
            continue
        if line == "TABLE:":
            rows, idx = _parse_table(lines, idx + 1)
            _add_table(doc, rows)
            continue
        if line.startswith("FIGURE:"):
            _add_figure(doc, root, line.removeprefix("FIGURE:").strip())
            idx += 1
            continue
        if line.startswith("> "):
            _add_callout(doc, line[2:])
            idx += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:])
            idx += 1
            continue

        paragraph_lines = [line]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].strip()
            if not nxt:
                break
            if nxt.startswith(("#", "- ", "> ", "FIGURE:", "TABLE:")):
                break
            paragraph_lines.append(nxt)
            idx += 1
        text = " ".join(paragraph_lines)
        style = "Reference" if in_references and text.startswith("[") else default_style
        p = doc.add_paragraph(style=style)
        _add_inline(p, text)


def build(source: Path, output: Path, root: Path) -> None:
    raw = source.read_text(encoding="utf-8")
    match = re.match(r"^---\n(.*?)\n---\n(.*)$", raw, flags=re.S)
    if not match:
        raise ValueError("Report source must start with YAML front matter")
    meta: dict[str, str] = {}
    for line in match.group(1).splitlines():
        key, separator, value = line.partition(":")
        if not separator:
            raise ValueError(f"Invalid front-matter line: {line}")
        meta[key.strip()] = value.strip()
    pages = [part.strip().splitlines() for part in match.group(2).split("---PAGE---")]

    doc = Document()
    _configure_document(doc, meta)
    _add_cover(doc, meta)
    _add_page_content(doc, root, pages[0], default_style="Executive Summary Body")
    doc.add_page_break()
    for page in pages[1:]:
        _add_page_content(doc, root, page)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--root", type=Path, default=Path.cwd())
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve(), args.root.resolve())


if __name__ == "__main__":
    main()
