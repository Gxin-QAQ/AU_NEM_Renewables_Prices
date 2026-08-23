"""Audit the frozen data, result, notebook and report contracts for Task 11."""

from __future__ import annotations

import argparse
import hashlib
import json
import site
import sys
from pathlib import Path
from typing import Any

import nbformat
import pandas as pd
import pyarrow.parquet as pq
from docx import Document


HISTORY_COUNT = 216
PILOT_COUNT = 47
REFERENCE_COUNT = 2
PANEL_ROWS = 263_040
FIVE_MINUTE_ROWS = 3_156_480
HEADLINE_ROWS = 210_399
DYNAMIC_ROWS = 209_928
WEEK_CLUSTERS = 314
SHARE_CAP = 2.7891799386467846
CAP_AFFECTED_ROWS = 211


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


class Audit:
    def __init__(self) -> None:
        self.checks: list[dict[str, Any]] = []

    def record(self, name: str, passed: bool, detail: Any) -> None:
        self.checks.append({"check": name, "passed": bool(passed), "detail": detail})

    def require(self, name: str, passed: bool, detail: Any) -> None:
        self.record(name, passed, detail)
        if not passed:
            raise AssertionError(f"{name}: {detail}")

    def result(self) -> dict[str, Any]:
        return {
            "status": "pass" if all(item["passed"] for item in self.checks) else "fail",
            "checks_run": len(self.checks),
            "checks": self.checks,
        }


def audit_environment(root: Path, audit: Audit) -> None:
    expected = (root / ".venv").resolve()
    audit.require("project .venv prefix", Path(sys.prefix).resolve() == expected, sys.prefix)
    audit.require("project .venv executable", expected in Path(sys.executable).absolute().parents, sys.executable)
    audit.require("user site disabled", site.ENABLE_USER_SITE is False, site.ENABLE_USER_SITE)


def audit_manifest(
    root: Path,
    audit: Audit,
    manifest_path: Path,
    frozen_path: Path,
    expected_rows: int,
    full_checksums: bool,
) -> pd.DataFrame:
    manifest = pd.read_csv(manifest_path)
    audit.require(f"{manifest_path.name} row count", len(manifest) == expected_rows, len(manifest))
    audit.require(
        f"{manifest_path.name} frozen copy",
        manifest_path.read_bytes() == frozen_path.read_bytes(),
        str(frozen_path),
    )
    failures: list[str] = []
    for row in manifest.itertuples(index=False):
        path = root / row.local_path
        if not path.exists():
            failures.append(f"missing:{row.local_path}")
            continue
        if path.stat().st_size != int(row.bytes):
            failures.append(f"size:{row.local_path}")
            continue
        if full_checksums and sha256(path) != row.sha256:
            failures.append(f"sha256:{row.local_path}")
    audit.require(
        f"{manifest_path.name} local integrity",
        not failures,
        {"files": len(manifest), "full_checksums": full_checksums, "failures": failures[:10]},
    )
    return manifest


def audit_sources(root: Path, audit: Audit, full_checksums: bool) -> None:
    history = audit_manifest(
        root,
        audit,
        root / "data/raw/history_manifest.csv",
        root / "provenance/aemo_history_manifest.csv",
        HISTORY_COUNT,
        full_checksums,
    )
    audit_manifest(
        root,
        audit,
        root / "data/raw/manifest.csv",
        root / "provenance/aemo_pilot_manifest.csv",
        PILOT_COUNT,
        full_checksums,
    )
    audit_manifest(
        root,
        audit,
        root / "data/external/reference_manifest.csv",
        root / "provenance/reference_manifest.csv",
        REFERENCE_COUNT,
        full_checksums,
    )
    periods = history.groupby("dataset")["period"].nunique().to_dict()
    audit.require("72 months per AEMO table family", periods == {"price": 72, "region": 72, "scada": 72}, periods)


def audit_partitions(root: Path, audit: Audit) -> None:
    expected = {
        "data/interim/history_generation_demand_5min": FIVE_MINUTE_ROWS,
        "data/interim/history_generation_demand_hour": PANEL_ROWS,
        "data/processed/nem_region_5min": FIVE_MINUTE_ROWS,
        "data/interim/history_price_hour": PANEL_ROWS,
    }
    for relative, expected_rows in expected.items():
        files = sorted((root / relative).glob("*.parquet"))
        rows = sum(pq.ParquetFile(path).metadata.num_rows for path in files)
        detail = {"files": len(files), "rows": rows, "first": files[0].stem if files else None, "last": files[-1].stem if files else None}
        audit.require(
            f"partition contract: {relative}",
            len(files) == 72 and rows == expected_rows and files[0].stem == "2019-07" and files[-1].stem == "2025-06",
            detail,
        )


def audit_panels(root: Path, audit: Audit) -> None:
    panel = pd.read_parquet(root / "data/processed/nem_region_hour.parquet")
    model = pd.read_parquet(root / "data/processed/nem_region_hour_model.parquet")
    audit.require("hourly panel rows", len(panel) == PANEL_ROWS, len(panel))
    audit.require("model frame rows", len(model) == PANEL_ROWS, len(model))
    audit.require("unique region-hour key", not panel.duplicated(["timestamp", "region"]).any(), "timestamp + region")
    audit.require("five NEM regions", set(panel["region"]) == {"NSW1", "VIC1", "QLD1", "SA1", "TAS1"}, sorted(panel["region"].unique()))
    audit.require("headline sample rows", int(model["headline_sample"].sum()) == HEADLINE_ROWS, int(model["headline_sample"].sum()))
    audit.require("dynamic sample rows", int(model["dynamic_sample"].sum()) == DYNAMIC_ROWS, int(model["dynamic_sample"].sum()))
    weeks = int(model.loc[model["headline_sample"], "aest_week"].nunique())
    audit.require("AEST week clusters", weeks == WEEK_CLUSTERS, weeks)
    headline = model.loc[model["headline_sample"]]
    cap = float(headline["renewable_share_ws"].quantile(0.999))
    affected = int((headline["renewable_share_ws"] > cap).sum())
    audit.require("frozen p99.9 cap", abs(cap - SHARE_CAP) < 1e-12 and affected == CAP_AFFECTED_ROWS, {"cap": cap, "affected": affected})


def value_for(table: pd.DataFrame, outcome: str) -> tuple[float, float]:
    row = table.loc[
        table["section"].eq("headline_fixed_effects")
        & table["outcome_or_contrast"].eq(outcome)
    ].iloc[0]
    return float(row["estimate"]), float(row["std_error"])


def audit_results(root: Path, audit: Audit) -> None:
    headline = pd.read_csv(root / "outputs/tables/task8_headline_results.csv")
    expected = {
        "Continuous price (asinh RRP)": (-0.2701236461017536, 0.0125084190815185),
        "Continuous price (AUD/MWh level)": (-11.80204941815006, 0.8679344094493564),
        "Any negative price (LPM, probability points)": (0.035768999779566, 0.001377483495032),
        "Intrahour volatility (AUD/MWh SD level)": (-5.025188722162409, 0.6837490563883706),
    }
    for outcome, target in expected.items():
        observed = value_for(headline, outcome)
        audit.require(
            f"headline result: {outcome}",
            all(abs(left - right) < 1e-12 for left, right in zip(observed, target)),
            {"estimate": observed[0], "std_error": observed[1]},
        )
    required_tables = {
        "task8_coefficients.csv",
        "task8_linear_contrasts.csv",
        "task8_nonlinear_negative_price_summary.csv",
        "task9_robustness_coefficients.csv",
        "task9_heterogeneity_differences_holm.csv",
        "task9_week_multiplier_bootstrap.csv",
    }
    present = {path.name for path in (root / "outputs/tables").glob("*.csv")}
    audit.require("required result tables", required_tables <= present, sorted(required_tables - present))


def audit_reader_artifacts(root: Path, audit: Audit) -> None:
    text_files = {
        "README.md": ["README_EN.md", "A$11.80/MWh", "3.58", "5.03", "条件相关性"],
        "README_EN.md": ["README.md", "A$11.80/MWh", "3.58", "5.03", "conditional associations"],
        "data/README.md": ["README_EN.md", "263,040", "210,399", "209,928", "2.78918"],
        "data/README_EN.md": ["README.md", "263,040", "210,399", "209,928", "2.78918"],
        "report/research_report.md": ["A$11.80/MWh", "3.58", "A$5.03/MWh", "conditional associations"],
    }
    for relative, needles in text_files.items():
        text = (root / relative).read_text(encoding="utf-8")
        missing = [needle for needle in needles if needle not in text]
        audit.require(f"reader text: {relative}", not missing, missing)

    for filename in [
        "fig1_price_and_renewable_share_trends.png",
        "fig5_regional_price_heterogeneity.png",
        "fig6_price_robustness.png",
    ]:
        output = root / "outputs/figures" / filename
        asset = root / "assets/readme" / filename
        audit.require(f"README asset sync: {filename}", output.exists() and asset.exists() and sha256(output) == sha256(asset), filename)

    for notebook_path in sorted((root / "notebooks").glob("*.ipynb")):
        notebook = nbformat.read(notebook_path, as_version=4)
        code_cells = [cell for cell in notebook.cells if cell.cell_type == "code"]
        errors = [output for cell in code_cells for output in cell.get("outputs", []) if output.output_type == "error"]
        executed = sum(cell.execution_count is not None for cell in code_cells)
        audit.require(
            f"executed notebook: {notebook_path.name}",
            bool(code_cells) and executed == len(code_cells) and not errors,
            {"executed": executed, "code_cells": len(code_cells), "errors": len(errors)},
        )

    docx_path = root / "report/AU_NEM_Renewables_Prices_Research_Report.docx"
    document = Document(docx_path)
    docx_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    missing = [needle for needle in ["A$11.80/MWh", "3.58", "A$5.03/MWh", "conditional associations"] if needle not in docx_text]
    audit.require("DOCX headline consistency", not missing, missing)


def run(root: Path, full_checksums: bool = False) -> dict[str, Any]:
    audit = Audit()
    audit_environment(root, audit)
    audit_sources(root, audit, full_checksums)
    audit_partitions(root, audit)
    audit_panels(root, audit)
    audit_results(root, audit)
    audit_reader_artifacts(root, audit)
    return audit.result()


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--root", type=Path, default=Path(__file__).resolve().parents[1])
    parser.add_argument("--full-checksums", action="store_true", help="Recompute SHA-256 for every local source file")
    args = parser.parse_args()
    try:
        result = run(args.root.resolve(), full_checksums=args.full_checksums)
    except AssertionError as error:
        print(json.dumps({"status": "fail", "error": str(error)}, indent=2, ensure_ascii=False))
        return 1
    print(json.dumps(result, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
